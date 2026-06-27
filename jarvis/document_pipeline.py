"""Document text extraction, caching, and Q&A for warranty/work docs and learning."""

from __future__ import annotations

import hashlib
import html.parser
import json
import logging
import os
import re
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

log = logging.getLogger("jarvis")

DOCUMENT_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown", ".html", ".htm"}
TEXT_EXTENSIONS = {".txt", ".md", ".markdown"}
HTML_EXTENSIONS = {".html", ".htm"}
DOCUMENTS_DIR = Path(__file__).resolve().parent.parent / "data" / "documents"
CACHE_DIR = DOCUMENTS_DIR / ".cache"

_MAX_CONTEXT_CHARS = int(os.getenv("JARVIS_DOC_CONTEXT_CHARS", "14000"))
_CHUNK_CHARS = int(os.getenv("JARVIS_DOC_CHUNK_CHARS", "3200"))
_TOP_CHUNKS = int(os.getenv("JARVIS_DOC_TOP_CHUNKS", "6"))


@dataclass
class ParsedDocument:
    path: str
    title: str
    pages: list[str] = field(default_factory=list)
    page_count: int = 0
    char_count: int = 0
    source: str = ""

    @property
    def excerpt(self) -> str:
        text = self.full_text.strip()
        if len(text) <= 500:
            return text
        return text[:500].rstrip() + "…"

    @property
    def full_text(self) -> str:
        return "\n\n".join(p for p in self.pages if p.strip())

    def to_dict(self) -> dict[str, Any]:
        return {
            "path": self.path,
            "title": self.title,
            "page_count": self.page_count,
            "char_count": self.char_count,
            "source": self.source,
            "excerpt": self.excerpt,
        }


def documents_dir() -> Path:
    DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
    CACHE_DIR.mkdir(parents=True, exist_ok=True)
    return DOCUMENTS_DIR


def is_document_path(path: str | Path) -> bool:
    return Path(path).suffix.lower() in DOCUMENT_EXTENSIONS


def document_attachment_action(message: str) -> str:
    """Route document attach: learn, ingest, summarize, query, vision_ocr, or info."""
    lower = (message or "").lower().strip()
    if not lower or lower in ("please analyze the attached file.", "analyze", "(attachment)"):
        return "summarize"
    if re.search(
        r"\b(learn from|study (?:this|the)|read and learn|memorize (?:this|the)|teach yourself from)\b",
        lower,
    ):
        return "learn"
    if re.search(r"\b(ingest|add to (?:my )?(?:document )?library|index (?:this|the) (?:doc|document|file))\b", lower):
        return "ingest"
    if re.search(
        r"\b(read (all )?text|ocr|scan(ned)?|what text|transcribe|describe (this )?(page|pdf|document))\b",
        lower,
    ):
        return "vision"
    if re.search(
        r"\b(html|screenshot|image|picture|photo)\b",
        lower,
    ) and re.search(r"\b(convert|to code|describe|what do you see)\b", lower):
        return "vision"
    if re.search(
        r"\b(summarize|summary|overview|tl;dr|key points|main points|executive summary)\b",
        lower,
    ):
        return "summarize"
    if re.search(
        r"\b(page count|how many pages|table of contents|metadata|document info)\b",
        lower,
    ):
        return "info"
    if re.search(
        r"\b(what|which|when|where|who|how|does|is there|are there|find|clause|section|warranty|coverage|expire)\b",
        lower,
    ) or "?" in lower:
        return "query"
    return "summarize"


def _cache_key(path: Path) -> str:
    stat = path.stat()
    raw = f"{path.resolve()}|{stat.st_mtime_ns}|{stat.st_size}"
    return hashlib.sha256(raw.encode()).hexdigest()[:16]


def _read_cache(path: Path) -> ParsedDocument | None:
    key = _cache_key(path)
    cache_file = CACHE_DIR / f"{key}.json"
    if not cache_file.exists():
        return None
    try:
        data = json.loads(cache_file.read_text(encoding="utf-8"))
        if data.get("path") != str(path.resolve()):
            return None
        return ParsedDocument(
            path=str(path),
            title=data.get("title") or path.name,
            pages=list(data.get("pages") or []),
            page_count=int(data.get("page_count") or 0),
            char_count=int(data.get("char_count") or 0),
            source=data.get("source") or "",
        )
    except (json.JSONDecodeError, OSError, ValueError):
        return None


def _write_cache(doc: ParsedDocument) -> None:
    CACHE_DIR.mkdir(parents=True, exist_ok=True)
    path = Path(doc.path)
    cache_file = CACHE_DIR / f"{_cache_key(path)}.json"
    payload = {
        "path": str(path.resolve()),
        "title": doc.title,
        "pages": doc.pages,
        "page_count": doc.page_count,
        "char_count": doc.char_count,
        "source": doc.source,
    }
    cache_file.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")


def _clean_page_text(text: str) -> str:
    text = re.sub(r"\r\n?", "\n", text or "")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_pdf(path: Path) -> ParsedDocument:
    pages: list[str] = []
    source = "pymupdf"
    try:
        import fitz  # pymupdf

        doc = fitz.open(str(path))
        try:
            for i in range(doc.page_count):
                pages.append(_clean_page_text(doc.load_page(i).get_text("text")))
        finally:
            doc.close()
    except ImportError:
        source = "pdftotext"
        pages = _extract_pdf_pdftotext(path)
    except Exception as exc:
        log.debug("pymupdf extract failed for %s: %s", path, exc)
        pages = _extract_pdf_pdftotext(path)
        source = "pdftotext"

    full = "\n\n".join(pages)
    return ParsedDocument(
        path=str(path),
        title=path.stem or path.name,
        pages=pages,
        page_count=len(pages),
        char_count=len(full),
        source=source,
    )


def _extract_pdf_pdftotext(path: Path) -> list[str]:
    with tempfile.TemporaryDirectory() as td:
        out = Path(td) / "out.txt"
        proc = subprocess.run(
            ["pdftotext", "-layout", str(path), str(out)],
            capture_output=True,
            timeout=120,
        )
        if proc.returncode != 0 or not out.exists():
            err = (proc.stderr or b"").decode(errors="ignore").strip()
            raise ValueError(err or "PDF text extract failed — install pymupdf or poppler-utils")
        text = out.read_text(encoding="utf-8", errors="ignore")
    chunks = re.split(r"\f", text)
    pages = [_clean_page_text(p) for p in chunks if p.strip()]
    if not pages and text.strip():
        pages = [_clean_page_text(text)]
    return pages


def _extract_docx(path: Path) -> ParsedDocument:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("Install python-docx for Word files: pip install python-docx") from exc

    doc = Document(str(path))
    paragraphs = [_clean_page_text(p.text) for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            paragraphs.append("\n".join(rows))
    full = "\n\n".join(paragraphs)
    return ParsedDocument(
        path=str(path),
        title=path.stem or path.name,
        pages=paragraphs or [""],
        page_count=max(1, len(paragraphs)),
        char_count=len(full),
        source="python-docx",
    )


class _HTMLTextExtractor(html.parser.HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []
        self._skip = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in ("script", "style", "noscript"):
            self._skip += 1
        elif tag in ("p", "br", "div", "li", "h1", "h2", "h3", "h4", "tr", "section", "article"):
            self._parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in ("script", "style", "noscript") and self._skip:
            self._skip -= 1

    def handle_data(self, data: str) -> None:
        if self._skip:
            return
        text = (data or "").strip()
        if text:
            self._parts.append(text + " ")

    def get_text(self) -> str:
        return _clean_page_text("".join(self._parts))


def html_to_text(html: str) -> str:
    parser = _HTMLTextExtractor()
    try:
        parser.feed(html or "")
        parser.close()
    except html.parser.HTMLParseError:
        pass
    return parser.get_text()


def _title_from_html(html: str) -> str:
    m = re.search(r"<title[^>]*>([^<]+)</title>", html or "", re.I | re.S)
    return _clean_page_text(m.group(1)) if m else ""


def _extract_text_file(path: Path) -> ParsedDocument:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    pages = [_clean_page_text(raw)] if raw.strip() else [""]
    return ParsedDocument(
        path=str(path),
        title=path.stem or path.name,
        pages=pages,
        page_count=1,
        char_count=len(raw),
        source="text",
    )


def _extract_html_file(path: Path) -> ParsedDocument:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    title = _title_from_html(raw) or path.stem or path.name
    text = html_to_text(raw)
    pages = [_clean_page_text(text)] if text.strip() else [""]
    return ParsedDocument(
        path=str(path),
        title=title,
        pages=pages,
        page_count=1,
        char_count=len(text),
        source="html",
    )


def parse_text_content(
    text: str,
    *,
    title: str = "document",
    source: str = "text",
    path: str = "",
) -> ParsedDocument:
    """Build a ParsedDocument from raw text (OCR output, pasted content, etc.)."""
    cleaned = _clean_page_text(text or "")
    pages = [cleaned] if cleaned else [""]
    return ParsedDocument(
        path=path or f"inline:{title}",
        title=(title or "document").strip(),
        pages=pages,
        page_count=1,
        char_count=len(cleaned),
        source=source,
    )


def parse_document(path: str | Path, *, use_cache: bool = True) -> ParsedDocument:
    """Extract text from supported document types."""
    documents_dir()
    p = Path(path).expanduser().resolve()
    if not p.exists():
        raise FileNotFoundError(f"Document not found: {p}")
    suffix = p.suffix.lower()
    if suffix not in DOCUMENT_EXTENSIONS:
        raise ValueError(f"Unsupported document type: {suffix}")

    if use_cache:
        cached = _read_cache(p)
        if cached:
            return cached

    if suffix == ".pdf":
        parsed = _extract_pdf(p)
    elif suffix == ".docx":
        parsed = _extract_docx(p)
    elif suffix in TEXT_EXTENSIONS:
        parsed = _extract_text_file(p)
    elif suffix in HTML_EXTENSIONS:
        parsed = _extract_html_file(p)
    else:
        raise ValueError(f"Unsupported document type: {suffix}")

    if use_cache:
        _write_cache(parsed)
    return parsed


def low_text_warning(doc: ParsedDocument) -> str | None:
    if doc.page_count <= 0:
        return "Could not read any pages from this document."
    if doc.char_count >= max(80, doc.page_count * 25):
        return None
    return (
        f"This document looks scanned or image-only ({doc.char_count} characters across "
        f"{doc.page_count} page(s)). For OCR, attach the file and ask to **read all text in this document** "
        "or **learn from this document with OCR**."
    )


def _chunk_document(doc: ParsedDocument) -> list[dict[str, Any]]:
    chunks: list[dict[str, Any]] = []
    buf = ""
    page_start = 1
    page_end = 1
    for idx, page in enumerate(doc.pages, start=1):
        if not page.strip():
            continue
        block = f"[Page {idx}]\n{page}"
        if buf and len(buf) + len(block) + 2 > _CHUNK_CHARS:
            chunks.append({"text": buf.strip(), "page_start": page_start, "page_end": page_end})
            buf = block
            page_start = idx
            page_end = idx
        else:
            buf = f"{buf}\n\n{block}".strip() if buf else block
            page_end = idx
    if buf.strip():
        chunks.append({"text": buf.strip(), "page_start": page_start, "page_end": page_end})
    return chunks


def _score_chunk(chunk: str, query: str) -> float:
    text = chunk.lower()
    words = [w for w in re.findall(r"[a-z0-9]{3,}", query.lower()) if len(w) > 2]
    if not words:
        return 0.0
    score = 0.0
    for word in words:
        if word in text:
            score += 2.0 + text.count(word) * 0.25
    return score


def select_context(doc: ParsedDocument, question: str = "", *, for_summary: bool = False) -> str:
    """Pick relevant chunks without stuffing the full document into the prompt."""
    chunks = _chunk_document(doc)
    if not chunks:
        return doc.full_text[:_MAX_CONTEXT_CHARS]

    if for_summary:
        selected = chunks[: min(len(chunks), _TOP_CHUNKS)]
    else:
        ranked = sorted(
            (( _score_chunk(c["text"], question), i, c) for i, c in enumerate(chunks)),
            key=lambda x: (-x[0], x[1]),
        )
        top = [c for score, _, c in ranked if score > 0][: _TOP_CHUNKS]
        selected = top or chunks[: min(3, len(chunks))]

    parts = []
    used = 0
    for chunk in selected:
        label = chunk["page_start"]
        if chunk["page_end"] != chunk["page_start"]:
            label = f"{chunk['page_start']}-{chunk['page_end']}"
        block = f"--- Pages {label} ---\n{chunk['text']}"
        if used + len(block) > _MAX_CONTEXT_CHARS:
            remain = _MAX_CONTEXT_CHARS - used
            if remain > 400:
                parts.append(block[:remain] + "…")
            break
        parts.append(block)
        used += len(block)
    return "\n\n".join(parts)


def format_document_info(doc: ParsedDocument) -> str:
    warn = low_text_warning(doc)
    lines = [
        f"**{doc.title}**",
        f"- Pages: {doc.page_count}",
        f"- Characters extracted: {doc.char_count:,}",
        f"- Extractor: {doc.source or 'unknown'}",
    ]
    if warn:
        lines.append(f"- Note: {warn}")
    elif doc.excerpt:
        lines.extend(["", "**Preview**", doc.excerpt])
    return "\n".join(lines)


def summarize_document(doc: ParsedDocument) -> str:
    from jarvis import llm

    warn = low_text_warning(doc)
    if warn and doc.char_count < 40:
        return warn

    context = select_context(doc, for_summary=True)
    system = (
        "You summarize documents for a busy user doing warranty and administration work. "
        "Use ONLY the provided excerpts. Structure with short sections and bullet points. "
        "Highlight dates, parties, coverage, exclusions, deadlines, and action items. "
        "If text is partial, say what is missing."
    )
    user = (
        f"Document: {doc.title}\n"
        f"Pages: {doc.page_count}\n\n"
        f"Excerpts:\n{context}\n\n"
        "Write a concise summary."
    )
    return llm.ask_with_system(llm.general_model(), system, user)


def answer_document(doc: ParsedDocument, question: str) -> str:
    from jarvis import llm

    warn = low_text_warning(doc)
    if warn and doc.char_count < 40:
        return warn

    context = select_context(doc, question=question, for_summary=False)
    system = (
        "You answer questions about documents using ONLY the provided excerpts. "
        "Quote or paraphrase accurately. Mention page numbers when visible in excerpts. "
        "If the answer is not in the excerpts, say so clearly — do not invent clauses or dates."
    )
    user = (
        f"Document: {doc.title}\n\n"
        f"Excerpts:\n{context}\n\n"
        f"Question: {question.strip()}"
    )
    return llm.ask_with_system(llm.general_model(), system, user)


def list_library_documents(limit: int = 50) -> list[dict[str, Any]]:
    root = documents_dir()
    items: list[dict[str, Any]] = []
    for path in sorted(root.rglob("*"), key=lambda p: p.stat().st_mtime, reverse=True):
        if not path.is_file() or path.suffix.lower() not in DOCUMENT_EXTENSIONS:
            continue
        if path.name.startswith(".") or ".cache" in path.parts:
            continue
        items.append({
            "name": path.name,
            "path": str(path),
            "size": path.stat().st_size,
            "modified": path.stat().st_mtime,
        })
        if len(items) >= limit:
            break
    return items
